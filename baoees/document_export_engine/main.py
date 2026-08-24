import json
from datetime import datetime
from pathlib import Path

from baoees.document_export_engine.libreoffice_bridge import (
    DocumentExportLibreOfficeBridge,
)


class DocumentExportEngine:

    def __init__(self):
        self.document_result = {}
        self.office_bridge = DocumentExportLibreOfficeBridge()

    def convert_office_document(self, input_path, target_format, output_dir):
        return self.office_bridge.convert_office_document(
            input_path,
            target_format,
            output_dir,
        )

    def open_office_document(self, input_path):
        return self.office_bridge.open_office_document(input_path)

    def libreoffice_capability(self):
        return self.office_bridge.libreoffice_capability()

    def create_documents(self, project_result=None, reporting_result=None, export_result=None):
        project_result = project_result or {}
        reporting_result = reporting_result or {}
        export_result = export_result or {}

        export_folder = export_result.get("export_folder")

        if not export_folder:
            self.document_result = {
                "engine": "DocumentExportEngine",
                "version": "1.1",
                "status": "DOCUMENT_EXPORT_MISLUKT",
                "reason": "Geen export_folder gevonden in export_result."
            }
            return self.document_result

        export_folder_path = Path(export_folder)
        export_folder_path.mkdir(parents=True, exist_ok=True)

        report_data = self.build_report_data(project_result, reporting_result)

        txt_path = export_folder_path / "projectrapport.txt"
        docx_path = export_folder_path / "projectrapport.docx"
        pdf_path = export_folder_path / "projectrapport.pdf"
        json_path = export_folder_path / "projectrapport_documentdata.json"

        txt_result = self.create_txt_report(txt_path, report_data)
        json_result = self.create_json_report(json_path, report_data)
        docx_result = self.create_docx_report(docx_path, report_data)
        pdf_result = self.create_pdf_report(pdf_path, report_data)

        self.document_result = {
            "engine": "DocumentExportEngine",
            "version": "1.1",
            "status": "DOCUMENT_EXPORT_GEREED",
            "export_folder": str(export_folder_path),
            "documents": [
                txt_result,
                json_result,
                docx_result,
                pdf_result
            ],
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "note": "Professionele basis DOCX/PDF-export aangemaakt met titelblad, hoofdstukken en metadata."
        }

        return self.document_result

    def build_report_data(self, project_result, reporting_result):
        project_name = project_result.get("project_name", "Onbekend project")

        report_title = reporting_result.get(
            "report_title",
            f"BAOEES Projectrapport - {project_name}"
        )

        sections = reporting_result.get("sections", [])

        return {
            "title": report_title,
            "subtitle": "Autonoom gegenereerd projectrapport",
            "system": "BAOEES V3.1",
            "engine": "DocumentExportEngine v1.1",
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "project": {
                "project_name": project_result.get("project_name", "Onbekend"),
                "project_type": project_result.get("project_type", "Onbekend"),
                "location": project_result.get("location", "Onbekend"),
                "country": project_result.get("country", "Onbekend"),
                "description": project_result.get("project_description", "")
            },
            "sections": sections,
            "disclaimer": (
                "Dit rapport is automatisch gegenereerd door BAOEES. "
                "De inhoud is geschikt als conceptbasis en moet bij formele indiening "
                "altijd worden gecontroleerd door een bevoegd deskundige."
            )
        }

    def create_txt_report(self, file_path, report_data):
        lines = []

        lines.append(report_data["title"])
        lines.append("=" * len(report_data["title"]))
        lines.append("")
        lines.append(report_data["subtitle"])
        lines.append("")
        lines.append(f"Systeem: {report_data['system']}")
        lines.append(f"Engine: {report_data['engine']}")
        lines.append(f"Datum: {report_data['created_at']}")
        lines.append("")

        project = report_data["project"]

        lines.append("PROJECTGEGEVENS")
        lines.append("---------------")
        lines.append(f"Projectnaam: {project['project_name']}")
        lines.append(f"Projecttype: {project['project_type']}")
        lines.append(f"Locatie: {project['location']}")
        lines.append(f"Land: {project['country']}")
        lines.append(f"Omschrijving: {project['description']}")
        lines.append("")

        lines.append("INHOUD")
        lines.append("------")

        for section in report_data["sections"]:
            chapter = section.get("chapter", "?")
            title = section.get("title", "Onbekend hoofdstuk")
            status = section.get("status", "CONCEPT")
            content_summary = section.get("content_summary", {})

            lines.append("")
            lines.append(f"{chapter}. {title}")
            lines.append("-" * 40)
            lines.append(f"Status: {status}")
            lines.append("")
            lines.append(json.dumps(content_summary, ensure_ascii=False, indent=2))

        lines.append("")
        lines.append("DISCLAIMER")
        lines.append("----------")
        lines.append(report_data["disclaimer"])

        with open(file_path, "w", encoding="utf-8") as file:
            file.write("\n".join(lines))

        return {
            "file": str(file_path),
            "format": "TXT",
            "status": "AANGEMAAKT"
        }

    def create_json_report(self, file_path, report_data):
        with open(file_path, "w", encoding="utf-8") as file:
            json.dump(report_data, file, ensure_ascii=False, indent=4)

        return {
            "file": str(file_path),
            "format": "JSON",
            "status": "AANGEMAAKT"
        }

    def create_docx_report(self, file_path, report_data):
        try:
            from docx import Document
            from docx.shared import Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            document = Document()

            styles = document.styles
            styles["Normal"].font.name = "Arial"
            styles["Normal"].font.size = Pt(10)

            title = document.add_paragraph()
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_run = title.add_run(report_data["title"])
            title_run.bold = True
            title_run.font.size = Pt(18)

            subtitle = document.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_run = subtitle.add_run(report_data["subtitle"])
            subtitle_run.italic = True
            subtitle_run.font.size = Pt(12)

            document.add_paragraph("")
            document.add_paragraph(f"Systeem: {report_data['system']}")
            document.add_paragraph(f"Engine: {report_data['engine']}")
            document.add_paragraph(f"Datum: {report_data['created_at']}")

            document.add_page_break()

            document.add_heading("1. Projectgegevens", level=1)

            project = report_data["project"]

            table = document.add_table(rows=0, cols=2)
            table.style = "Table Grid"

            self.add_docx_table_row(table, "Projectnaam", project["project_name"])
            self.add_docx_table_row(table, "Projecttype", project["project_type"])
            self.add_docx_table_row(table, "Locatie", project["location"])
            self.add_docx_table_row(table, "Land", project["country"])
            self.add_docx_table_row(table, "Omschrijving", project["description"])

            document.add_heading("2. Rapporthoofdstukken", level=1)

            for section in report_data["sections"]:
                chapter = section.get("chapter", "?")
                section_title = section.get("title", "Onbekend hoofdstuk")
                status = section.get("status", "CONCEPT")
                content_summary = section.get("content_summary", {})

                document.add_heading(f"{chapter}. {section_title}", level=2)
                document.add_paragraph(f"Status: {status}")

                if isinstance(content_summary, dict):
                    for key, value in content_summary.items():
                        paragraph = document.add_paragraph()
                        paragraph.add_run(f"{key}: ").bold = True
                        paragraph.add_run(str(value))
                else:
                    document.add_paragraph(str(content_summary))

            document.add_heading("Disclaimer", level=1)
            document.add_paragraph(report_data["disclaimer"])

            document.save(file_path)

            return {
                "file": str(file_path),
                "format": "DOCX",
                "status": "AANGEMAAKT"
            }

        except Exception as error:
            fallback_path = Path(str(file_path) + ".fallback.txt")
            with open(fallback_path, "w", encoding="utf-8") as file:
                file.write("DOCX-export mislukt.\n")
                file.write(f"Foutmelding: {error}\n")
                file.write(json.dumps(report_data, ensure_ascii=False, indent=2))

            return {
                "file": str(fallback_path),
                "format": "DOCX_FALLBACK_TXT",
                "status": "FALLBACK_AANGEMAAKT",
                "error": str(error)
            }

    def add_docx_table_row(self, table, label, value):
        row = table.add_row()
        row.cells[0].text = str(label)
        row.cells[1].text = str(value)

    def create_pdf_report(self, file_path, report_data):
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.units import cm
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.platypus import (
                SimpleDocTemplate,
                Paragraph,
                Spacer,
                Table,
                TableStyle,
                PageBreak
            )
            from reportlab.lib import colors

            document = SimpleDocTemplate(
                str(file_path),
                pagesize=A4,
                rightMargin=2 * cm,
                leftMargin=2 * cm,
                topMargin=2 * cm,
                bottomMargin=2 * cm
            )

            styles = getSampleStyleSheet()

            title_style = ParagraphStyle(
                "BAOEES_Title",
                parent=styles["Title"],
                fontSize=18,
                leading=22,
                alignment=1,
                spaceAfter=20
            )

            heading_style = ParagraphStyle(
                "BAOEES_Heading",
                parent=styles["Heading1"],
                fontSize=14,
                leading=18,
                spaceBefore=12,
                spaceAfter=10
            )

            subheading_style = ParagraphStyle(
                "BAOEES_SubHeading",
                parent=styles["Heading2"],
                fontSize=11,
                leading=14,
                spaceBefore=10,
                spaceAfter=6
            )

            normal_style = ParagraphStyle(
                "BAOEES_Normal",
                parent=styles["Normal"],
                fontSize=9,
                leading=12
            )

            story = []

            story.append(Paragraph(report_data["title"], title_style))
            story.append(Paragraph(report_data["subtitle"], normal_style))
            story.append(Spacer(1, 0.5 * cm))
            story.append(Paragraph(f"Systeem: {report_data['system']}", normal_style))
            story.append(Paragraph(f"Engine: {report_data['engine']}", normal_style))
            story.append(Paragraph(f"Datum: {report_data['created_at']}", normal_style))
            story.append(PageBreak())

            story.append(Paragraph("1. Projectgegevens", heading_style))

            project = report_data["project"]

            project_table_data = [
                ["Projectnaam", project["project_name"]],
                ["Projecttype", project["project_type"]],
                ["Locatie", project["location"]],
                ["Land", project["country"]],
                ["Omschrijving", project["description"]]
            ]

            project_table = Table(project_table_data, colWidths=[4 * cm, 11 * cm])
            project_table.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("BACKGROUND", (0, 0), (0, -1), colors.lightgrey),
                ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5)
            ]))

            story.append(project_table)
            story.append(Spacer(1, 0.5 * cm))

            story.append(Paragraph("2. Rapporthoofdstukken", heading_style))

            for section in report_data["sections"]:
                chapter = section.get("chapter", "?")
                section_title = section.get("title", "Onbekend hoofdstuk")
                status = section.get("status", "CONCEPT")
                content_summary = section.get("content_summary", {})

                story.append(Paragraph(f"{chapter}. {section_title}", subheading_style))
                story.append(Paragraph(f"Status: {status}", normal_style))

                if isinstance(content_summary, dict):
                    for key, value in content_summary.items():
                        safe_text = f"<b>{key}:</b> {value}"
                        story.append(Paragraph(str(safe_text), normal_style))
                else:
                    story.append(Paragraph(str(content_summary), normal_style))

                story.append(Spacer(1, 0.25 * cm))

            story.append(Paragraph("Disclaimer", heading_style))
            story.append(Paragraph(report_data["disclaimer"], normal_style))

            document.build(story)

            return {
                "file": str(file_path),
                "format": "PDF",
                "status": "AANGEMAAKT"
            }

        except Exception as error:
            fallback_path = Path(str(file_path) + ".fallback.txt")
            with open(fallback_path, "w", encoding="utf-8") as file:
                file.write("PDF-export mislukt.\n")
                file.write(f"Foutmelding: {error}\n")
                file.write(json.dumps(report_data, ensure_ascii=False, indent=2))

            return {
                "file": str(fallback_path),
                "format": "PDF_FALLBACK_TXT",
                "status": "FALLBACK_AANGEMAAKT",
                "error": str(error)
            }

    def get_document_result(self):
        return self.document_result

    def run(self):
        print("PDF/DOCX Export Engine actief")