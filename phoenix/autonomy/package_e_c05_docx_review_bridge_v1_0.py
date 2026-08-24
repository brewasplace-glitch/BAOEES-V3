"""Project Phoenix Package-E C05 DOCX Review Bridge v1.0.

Thin orchestration layer:
existing Package-E evidence -> existing C05 validation-list contract ->
reviewer-friendly DOCX -> returned DOCX -> canonical reviewed-input JSON.

It does not perform professional review, invent acceptance criteria, or replace
the existing Package-E validator / professional review gates.
"""

from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from hashlib import sha256
import json
from pathlib import Path
from typing import Any

from docx import Document

from phoenix.autonomy.generated_input_validation_list_v1_0 import (
    build_validation_list,
)
from phoenix.engines.adapters.libreoffice_document_router_v1_0 import (
    create_pdf_companion,
)

SCHEMA_VERSION = "phoenix.package-e-c05-docx-review-bridge/1.0"
RETURN_SCHEMA_VERSION = "phoenix.package-e-c05-reviewed-input/1.0"
PACKAGE_ID = "PKG-E-ALTERNATE-PATH-INDEPENDENT-EVIDENCE"

ALLOWED_ACTIONS = {"CONFIRM", "MODIFY", "NOT_APPLICABLE", "DEFER"}

PACKAGE_E_FIELD_GROUPS = {
    "applicability": "project_decision",
    "methodology_accepted": "project_decision",
    "methodology_acceptance_reference": "project_decision",
    "primary_source_record_id": "project_decision",
    "minimum_residual_capacity_proxy_ratio": "project_decision",
    "evidence_origin": "independent_engineering_evidence",
    "evidence_reference": "independent_engineering_evidence",
    "evidence_file_name": "independent_engineering_evidence",
    "evidence_sha256": "independent_engineering_evidence",
    "independently_verified_alternate_path": "independent_engineering_evidence",
    "review_status": "independent_review",
    "review_reference": "independent_review",
    "reviewer_identity_reference": "independent_review",
    "reviewer_qualification_reference": "independent_review",
    "review_scope": "independent_review",
    "review_date": "independent_review",
    "review_comments": "independent_review",
}


class ReviewBridgeError(ValueError):
    """Fail-closed error for malformed review bridge data."""


def _read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8-sig"))


def _write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(
        json.dumps(payload, indent=2, ensure_ascii=False) + "\n",
        encoding="utf-8",
    )


def _sha256(path: Path) -> str:
    h = sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            h.update(block)
    return h.hexdigest()


def _as_json_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value
    return json.dumps(value, ensure_ascii=False)


def _parse_review_value(text: str) -> Any:
    text = text.strip()
    if not text:
        return None
    lowered = text.lower()
    if lowered == "true":
        return True
    if lowered == "false":
        return False
    if lowered in {"null", "none"}:
        return None
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        return text


def _runtime_paths(repo: Path, project_id: str) -> tuple[Path, Path]:
    root = (
        repo / "projects" / "runtime" / project_id / "results" /
        "session_adapters" / "structural_engineering" /
        "validated_v8_1_to_v8_12" / "v8_6"
    )
    return (
        root / "r9_3_residual_capacity_stability_design_basis.json",
        root / "r9_5_project_stability_design_basis_decision.json",
    )


def build_package_e_candidate_inputs(
    repo_path: str | Path,
    project_id: str,
) -> dict[str, Any]:
    repo = Path(repo_path)
    r93_path, r95_path = _runtime_paths(repo, project_id)
    if not r93_path.is_file():
        raise ReviewBridgeError(f"R9.3 evidence missing: {r93_path}")
    if not r95_path.is_file():
        raise ReviewBridgeError(f"R9.5 evidence missing: {r95_path}")

    r93 = _read_json(r93_path)
    r95 = _read_json(r95_path)

    alt = (
        (r95.get("decision_register") or {})
        .get("ALTERNATE_LOAD_PATH_EVIDENCE")
    )
    if not isinstance(alt, dict):
        raise ReviewBridgeError("R9.5 alternate-load-path decision record missing")

    cases = (
        (r93.get("alternate_path_capacity_screening") or {})
        .get("cases") or []
    )
    ratios = [
        float(case["governing_residual_capacity_proxy_ratio"])
        for case in cases
        if isinstance(case, dict)
        and case.get("governing_residual_capacity_proxy_ratio") is not None
    ]
    if not ratios:
        raise ReviewBridgeError("R9.3 contains no governing proxy ratios")

    source_r93 = r93_path.relative_to(repo).as_posix()
    source_r95 = r95_path.relative_to(repo).as_posix()

    generated_inputs = [
        {
            "field": "project_id",
            "value": project_id,
            "classification": "AUTO_DERIVED",
            "source": source_r95,
            "confidence": "HIGH",
            "reason": "Active Phoenix runtime project identifier.",
        },
        {
            "field": "package_id",
            "value": PACKAGE_ID,
            "classification": "AUTO_DERIVED",
            "source": "Phoenix Package-E contract",
            "confidence": "HIGH",
            "reason": "Existing Package-E capability identifier.",
        },
        {
            "field": "current_r9_5_state",
            "value": alt.get("state"),
            "classification": "AUTO_DERIVED",
            "source": source_r95,
            "confidence": "HIGH",
            "reason": "Current R9.5 alternate-path decision state.",
        },
        {
            "field": "r9_3_screening_case_count",
            "value": len(ratios),
            "classification": "AUTO_DERIVED",
            "source": source_r93,
            "confidence": "HIGH",
            "reason": "Count of existing Phoenix alternate-path screening cases.",
        },
        {
            "field": "r9_3_observed_min_proxy_ratio",
            "value": min(ratios),
            "classification": "AUTO_DERIVED",
            "source": source_r93,
            "confidence": "HIGH",
            "reason": (
                "Observed Phoenix screening result only; this is NOT automatically "
                "the professional acceptance criterion."
            ),
        },
        {
            "field": "r9_3_observed_max_proxy_ratio",
            "value": max(ratios),
            "classification": "AUTO_DERIVED",
            "source": source_r93,
            "confidence": "HIGH",
            "reason": "Observed Phoenix screening context.",
        },
        {
            "field": "applicability",
            "value": None,
            "classification": "PROFESSIONAL_REVIEW_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": "Explicit Package-E applicability decision is required.",
        },
        {
            "field": "methodology_accepted",
            "value": None,
            "classification": "PROFESSIONAL_REVIEW_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": "Phoenix may not claim methodology acceptance.",
        },
        {
            "field": "methodology_acceptance_reference",
            "value": None,
            "classification": "HUMAN_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": "Traceable reference to the actual methodology decision.",
        },
        {
            "field": "primary_source_record_id",
            "value": None,
            "classification": "HUMAN_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": "Qualified source-record identifier required by Package E.",
        },
        {
            "field": "minimum_residual_capacity_proxy_ratio",
            "value": None,
            "classification": "PROFESSIONAL_REVIEW_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": (
                "Acceptance criterion must be independently supplied/accepted. "
                "The R9.3 observed minimum is context only."
            ),
        },
        {
            "field": "evidence_origin",
            "value": None,
            "classification": "HUMAN_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": "Independent evidence origin must be explicitly supplied.",
        },
        {
            "field": "evidence_reference",
            "value": None,
            "classification": "HUMAN_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": "Reference to independent engineering evidence.",
        },
        {
            "field": "evidence_file_name",
            "value": None,
            "classification": "HUMAN_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": "Filename of the independently supplied evidence.",
        },
        {
            "field": "evidence_sha256",
            "value": None,
            "classification": "HUMAN_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": "SHA-256 of the independently supplied evidence file.",
        },
        {
            "field": "independently_verified_alternate_path",
            "value": None,
            "classification": "PROFESSIONAL_REVIEW_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": "Must reflect an actual independent engineering verification.",
        },
        {
            "field": "review_status",
            "value": None,
            "classification": "PROFESSIONAL_REVIEW_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": "Actual review status; Phoenix cannot fabricate REVIEWED.",
        },
        {
            "field": "review_reference",
            "value": None,
            "classification": "HUMAN_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": "Reference to the actual independent review.",
        },
        {
            "field": "reviewer_identity_reference",
            "value": None,
            "classification": "HUMAN_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": "Reviewer identity/reference must come from the reviewer.",
        },
        {
            "field": "reviewer_qualification_reference",
            "value": None,
            "classification": "HUMAN_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": "Reviewer qualification reference must be explicit.",
        },
        {
            "field": "review_scope",
            "value": None,
            "classification": "PROFESSIONAL_REVIEW_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": "Scope of actual independent review.",
        },
        {
            "field": "review_date",
            "value": None,
            "classification": "HUMAN_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": "Date of actual review.",
        },
        {
            "field": "review_comments",
            "value": None,
            "classification": "HUMAN_REQUIRED",
            "source": source_r95,
            "confidence": None,
            "reason": "Optional reviewer comments or qualifications.",
        },
    ]

    return {
        "schema_version": SCHEMA_VERSION,
        "project_id": project_id,
        "package_id": PACKAGE_ID,
        "source_files": {
            "r9_3": source_r93,
            "r9_5": source_r95,
        },
        "generated_inputs": generated_inputs,
        "safety": {
            "r9_3_screening_is_independent_evidence": False,
            "automatic_professional_review_claim": False,
            "automatic_acceptance_criterion_invention": False,
            "production_release": "LOCKED",
        },
    }


def create_review_docx(
    validation_list: dict[str, Any],
    output_path: str | Path,
) -> Path:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("PROJECT PHOENIX — Package E Input Validation", level=1)
    doc.add_paragraph(
        "Phoenix heeft onderstaande input zoveel mogelijk zelf gegenereerd. "
        "Controleer de waarden en kies per regel een actie: CONFIRM, MODIFY, "
        "NOT_APPLICABLE of DEFER. Alleen de kolommen REVIEWER ACTION, "
        "REVIEWER VALUE en REVIEWER COMMENT hoeven te worden aangepast."
    )
    doc.add_paragraph(
        "Belangrijk: automatisch afgeleide R9.3-resultaten zijn context en vormen "
        "geen onafhankelijke professionele review of automatisch acceptatiecriterium."
    )
    doc.add_paragraph(
        f"Project: {validation_list.get('project_id')} | "
        f"Package: {validation_list.get('package_id')} | "
        "Status: CONCEPT_INPUT_VALIDATION_REQUIRED"
    )

    table = doc.add_table(rows=1, cols=10)
    table.style = "Table Grid"
    headers = [
        "PHX FIELD ID",
        "FIELD",
        "PHOENIX VALUE",
        "CLASSIFICATION",
        "SOURCE",
        "CONFIDENCE",
        "RATIONALE",
        "REVIEWER ACTION",
        "REVIEWER VALUE",
        "REVIEWER COMMENT",
    ]
    for idx, label in enumerate(headers):
        table.rows[0].cells[idx].text = label

    for item in validation_list["items"]:
        row = table.add_row().cells
        row[0].text = str(item["phoenix_field_id"])
        row[1].text = str(item["field"])
        row[2].text = _as_json_text(item["phoenix_value"])
        row[3].text = str(item["classification"])
        row[4].text = item.get("source") or ""
        row[5].text = item.get("confidence") or ""
        row[6].text = item.get("rationale") or ""
        row[7].text = ""
        row[8].text = ""
        row[9].text = ""

    doc.add_heading("Reviewer declaration", level=2)
    doc.add_paragraph(
        "Door REVIEWER ACTIONS in te vullen bevestigt de reviewer uitsluitend "
        "de expliciet gekozen bevestigingen/correcties. Phoenix leidt hieruit "
        "geen professionele goedkeuring af buiten de daadwerkelijk ingevulde velden."
    )
    doc.add_paragraph(
        "Lever het ingevulde DOCX samen met eventueel onafhankelijk evidencebestand "
        "terug aan Phoenix."
    )

    doc.save(output)
    return output


def prepare_package_e_review(
    repo_path: str | Path,
    project_id: str,
    output_dir: str | Path,
) -> dict[str, str]:
    output = Path(output_dir)
    output.mkdir(parents=True, exist_ok=True)

    candidates = build_package_e_candidate_inputs(repo_path, project_id)
    candidate_path = output / "PHOENIX_PACKAGE_E_AUTO_GENERATED_INPUTS_CONCEPT.json"
    _write_json(candidate_path, candidates)

    validation = build_validation_list(
        candidates["generated_inputs"],
        project_id=project_id,
        package_id=PACKAGE_ID,
    )
    validation_path = output / "PHOENIX_PACKAGE_E_GENERATED_INPUT_VALIDATION_LIST.json"
    _write_json(validation_path, validation)

    docx_path = output / "PHOENIX_PACKAGE_E_INPUT_VALIDATION_REVIEW_FORM.docx"
    create_review_docx(validation, docx_path)

    review_pdf_path = None
    review_pdf_conversion = {
        "status": "NOT_CREATED",
        "engine": "LibreOffice",
        "error": None,
    }
    try:
        review_pdf_conversion = create_pdf_companion(docx_path, output)
        candidate_pdf = Path(review_pdf_conversion["output"])
        if candidate_pdf.is_file() and candidate_pdf.stat().st_size > 0:
            review_pdf_path = candidate_pdf
        else:
            review_pdf_conversion = {
                **review_pdf_conversion,
                "status": "FAILED_OUTPUT_VALIDATION",
                "error": "LibreOffice did not produce a non-empty PDF companion.",
            }
    except Exception as exc:
        review_pdf_conversion = {
            "status": "FAILED_NONFATAL",
            "engine": "LibreOffice",
            "error": str(exc),
        }

    manifest_files = [
        {"name": candidate_path.name, "sha256": _sha256(candidate_path)},
        {"name": validation_path.name, "sha256": _sha256(validation_path)},
        {"name": docx_path.name, "sha256": _sha256(docx_path)},
    ]
    if review_pdf_path is not None:
        manifest_files.append({
            "name": review_pdf_path.name,
            "sha256": _sha256(review_pdf_path),
        })

    manifest = {
        "schema_version": "phoenix.package-e-c05-review-manifest/1.0",
        "project_id": project_id,
        "package_id": PACKAGE_ID,
        "files": manifest_files,
        "review_pdf_companion": review_pdf_conversion,
        "next_action": "REVIEWER_VALIDATES_OR_CORRECTS_DOCX_AND_RETURNS_IT",
        "production_release": "LOCKED",
    }
    manifest_path = output / "PHOENIX_PACKAGE_E_C05_REVIEW_MANIFEST.json"
    _write_json(manifest_path, manifest)

    return {
        "candidate_json": str(candidate_path),
        "validation_json": str(validation_path),
        "review_docx": str(docx_path),
        "review_pdf": str(review_pdf_path) if review_pdf_path is not None else "",
        "manifest_json": str(manifest_path),
    }


def ingest_review_docx(
    review_docx: str | Path,
    validation_json: str | Path,
    output_json: str | Path,
    evidence_dir: str | Path | None = None,
) -> dict[str, Any]:
    docx_path = Path(review_docx)
    validation_path = Path(validation_json)
    output_path = Path(output_json)

    if not docx_path.is_file():
        raise ReviewBridgeError(f"review DOCX not found: {docx_path}")
    if not validation_path.is_file():
        raise ReviewBridgeError(f"validation JSON not found: {validation_path}")

    validation = _read_json(validation_path)
    by_id = {
        item["phoenix_field_id"]: item
        for item in validation.get("items", [])
    }
    if not by_id:
        raise ReviewBridgeError("validation list has no items")

    doc = Document(docx_path)
    if not doc.tables:
        raise ReviewBridgeError("review DOCX contains no validation table")

    table = doc.tables[0]
    reviewed_items: list[dict[str, Any]] = []
    seen_ids: set[str] = set()

    for row in table.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) < 10:
            raise ReviewBridgeError("review table row has fewer than 10 cells")

        field_id = cells[0]
        if field_id not in by_id:
            raise ReviewBridgeError(f"unknown PHX field id in returned DOCX: {field_id}")
        if field_id in seen_ids:
            raise ReviewBridgeError(f"duplicate PHX field id in returned DOCX: {field_id}")
        seen_ids.add(field_id)

        original = by_id[field_id]
        action = cells[7].upper()
        reviewer_value_text = cells[8]
        reviewer_comment = cells[9] or None

        if not action:
            action = "DEFER"
        if action not in ALLOWED_ACTIONS:
            raise ReviewBridgeError(
                f"unsupported reviewer action {action!r} for {field_id}"
            )

        if action == "CONFIRM":
            reviewed_value = original.get("phoenix_value")
        elif action == "MODIFY":
            if not reviewer_value_text:
                raise ReviewBridgeError(
                    f"MODIFY requires REVIEWER VALUE for {field_id}"
                )
            reviewed_value = _parse_review_value(reviewer_value_text)
        elif action == "NOT_APPLICABLE":
            reviewed_value = None
        else:
            reviewed_value = original.get("phoenix_value")

        reviewed_items.append({
            "phoenix_field_id": field_id,
            "field": original["field"],
            "classification": original["classification"],
            "original_phoenix_value": original.get("phoenix_value"),
            "reviewer_action": action,
            "reviewed_value": reviewed_value,
            "reviewer_comment": reviewer_comment,
            "source": original.get("source"),
            "confidence": original.get("confidence"),
        })

    missing_rows = sorted(set(by_id) - seen_ids)
    if missing_rows:
        raise ReviewBridgeError(
            "returned DOCX omitted validation rows: " + ", ".join(missing_rows)
        )

    package_e = {
        "project_decision": {},
        "independent_engineering_evidence": {},
        "independent_review": {},
    }

    for item in reviewed_items:
        field = item["field"]
        group = PACKAGE_E_FIELD_GROUPS.get(field)
        if group:
            package_e[group][field] = item["reviewed_value"]

    evidence_validation = {
        "evidence_file_checked": False,
        "evidence_hash_matches": None,
        "resolved_evidence_file": None,
    }
    evidence_file_name = package_e["independent_engineering_evidence"].get(
        "evidence_file_name"
    )
    declared_hash = package_e["independent_engineering_evidence"].get(
        "evidence_sha256"
    )

    if evidence_file_name and evidence_dir:
        evidence_file = Path(evidence_dir) / str(evidence_file_name)
        if not evidence_file.is_file():
            raise ReviewBridgeError(
                f"declared evidence file not found: {evidence_file}"
            )
        actual_hash = _sha256(evidence_file)
        evidence_validation = {
            "evidence_file_checked": True,
            "evidence_hash_matches": (
                bool(declared_hash)
                and actual_hash.lower() == str(declared_hash).lower()
            ),
            "resolved_evidence_file": str(evidence_file),
            "actual_sha256": actual_hash,
        }
        if declared_hash and not evidence_validation["evidence_hash_matches"]:
            raise ReviewBridgeError("independent evidence SHA-256 mismatch")

    unresolved = [
        item["field"]
        for item in reviewed_items
        if item["reviewer_action"] == "DEFER"
        and item["classification"] in {
            "HUMAN_REQUIRED",
            "PROFESSIONAL_REVIEW_REQUIRED",
        }
    ]

    result = {
        "schema_version": RETURN_SCHEMA_VERSION,
        "project_id": validation.get("project_id"),
        "package_id": validation.get("package_id"),
        "source_review_docx": {
            "path": str(docx_path),
            "sha256": _sha256(docx_path),
        },
        "source_validation_json": {
            "path": str(validation_path),
            "sha256": _sha256(validation_path),
        },
        "reviewed_items": reviewed_items,
        "package_e_review_return": package_e,
        "evidence_validation": evidence_validation,
        "unresolved_required_review_fields": unresolved,
        "ready_for_existing_package_e_validation": not unresolved,
        "safety": {
            "professional_review_fabricated": False,
            "automatic_approval": False,
            "production_release": "LOCKED",
        },
    }

    _write_json(output_path, result)
    return result
