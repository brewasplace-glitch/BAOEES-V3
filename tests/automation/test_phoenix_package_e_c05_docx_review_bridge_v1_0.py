import json
from pathlib import Path
import tempfile
import unittest

from docx import Document

from phoenix.autonomy.package_e_c05_docx_review_bridge_v1_0 import (
    ReviewBridgeError,
    build_package_e_candidate_inputs,
    ingest_review_docx,
    prepare_package_e_review,
)


class PackageEC05DocxReviewBridgeTests(unittest.TestCase):
    def _make_repo(self, root: Path) -> Path:
        repo = root / "repo"
        v86 = (
            repo / "projects" / "runtime" / "PHOENIX-PAT-001" /
            "results" / "session_adapters" / "structural_engineering" /
            "validated_v8_1_to_v8_12" / "v8_6"
        )
        v86.mkdir(parents=True)

        r93 = {
            "alternate_path_capacity_screening": {
                "cases": [
                    {"governing_residual_capacity_proxy_ratio": 0.9375},
                    {"governing_residual_capacity_proxy_ratio": 1.0},
                ]
            }
        }
        r95 = {
            "decision_register": {
                "ALTERNATE_LOAD_PATH_EVIDENCE": {
                    "state": "DECISION_OR_SOURCE_INPUT_REQUIRED",
                    "missing_requirements": [
                        "explicit_applicability_decision",
                        "independent_review_status_REVIEWED",
                    ],
                }
            }
        }
        (v86 / "r9_3_residual_capacity_stability_design_basis.json").write_text(
            json.dumps(r93), encoding="utf-8"
        )
        (v86 / "r9_5_project_stability_design_basis_decision.json").write_text(
            json.dumps(r95), encoding="utf-8"
        )
        return repo

    def test_candidate_inputs_use_r93_context_without_promoting_it(self):
        with tempfile.TemporaryDirectory() as td:
            repo = self._make_repo(Path(td))
            payload = build_package_e_candidate_inputs(repo, "PHOENIX-PAT-001")
            values = {x["field"]: x for x in payload["generated_inputs"]}
            self.assertEqual(values["r9_3_screening_case_count"]["value"], 2)
            self.assertEqual(values["r9_3_observed_min_proxy_ratio"]["value"], 0.9375)
            self.assertIsNone(values["minimum_residual_capacity_proxy_ratio"]["value"])
            self.assertEqual(
                values["minimum_residual_capacity_proxy_ratio"]["classification"],
                "PROFESSIONAL_REVIEW_REQUIRED",
            )
            self.assertFalse(
                payload["safety"]["r9_3_screening_is_independent_evidence"]
            )

    def test_prepare_generates_json_docx_and_manifest(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            repo = self._make_repo(root)
            out = root / "out"
            result = prepare_package_e_review(repo, "PHOENIX-PAT-001", out)
            for path in result.values():
                self.assertTrue(Path(path).is_file())
            doc = Document(result["review_docx"])
            self.assertEqual(len(doc.tables), 1)
            self.assertGreater(len(doc.tables[0].rows), 5)

    def test_returned_docx_can_be_ingested(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            repo = self._make_repo(root)
            out = root / "out"
            result = prepare_package_e_review(repo, "PHOENIX-PAT-001", out)

            doc = Document(result["review_docx"])
            for row in doc.tables[0].rows[1:]:
                field = row.cells[1].text.strip()
                if field in {"applicability", "methodology_accepted"}:
                    row.cells[7].text = "MODIFY"
                    row.cells[8].text = (
                        "APPLICABLE" if field == "applicability" else "true"
                    )
                elif field.startswith("r9_3_") or field in {
                    "project_id", "package_id", "current_r9_5_state"
                }:
                    row.cells[7].text = "CONFIRM"
                else:
                    row.cells[7].text = "DEFER"
            doc.save(result["review_docx"])

            review_json = out / "reviewed.json"
            reviewed = ingest_review_docx(
                result["review_docx"],
                result["validation_json"],
                review_json,
            )
            self.assertTrue(review_json.is_file())
            self.assertEqual(
                reviewed["package_e_review_return"]["project_decision"]["applicability"],
                "APPLICABLE",
            )
            self.assertTrue(
                reviewed["package_e_review_return"]["project_decision"]["methodology_accepted"]
            )
            self.assertFalse(reviewed["ready_for_existing_package_e_validation"])

    def test_unknown_field_id_fails_closed(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            repo = self._make_repo(root)
            out = root / "out"
            result = prepare_package_e_review(repo, "PHOENIX-PAT-001", out)
            doc = Document(result["review_docx"])
            doc.tables[0].rows[1].cells[0].text = "PHX-VAL-UNKNOWN"
            doc.save(result["review_docx"])
            with self.assertRaises(ReviewBridgeError):
                ingest_review_docx(
                    result["review_docx"],
                    result["validation_json"],
                    out / "bad.json",
                )

    def test_modify_without_value_fails_closed(self):
        with tempfile.TemporaryDirectory() as td:
            root = Path(td)
            repo = self._make_repo(root)
            out = root / "out"
            result = prepare_package_e_review(repo, "PHOENIX-PAT-001", out)
            doc = Document(result["review_docx"])
            doc.tables[0].rows[1].cells[7].text = "MODIFY"
            doc.tables[0].rows[1].cells[8].text = ""
            doc.save(result["review_docx"])
            with self.assertRaises(ReviewBridgeError):
                ingest_review_docx(
                    result["review_docx"],
                    result["validation_json"],
                    out / "bad.json",
                )


if __name__ == "__main__":
    unittest.main()
